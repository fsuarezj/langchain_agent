import pandas as pd
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from langchain_agent.assistants.helping_features.files_manager import FilesManager
from langchain_community.document_loaders import UnstructuredWordDocumentLoader

file_path = "tests/tmp/test_form.docx"
def create_test_input(path):
    input = FilesManager()
    input.load_file(path)
    with open("langchain_agent/assistants/examples/form_parser2.in", "w", encoding='utf-8') as file:
        file.write(input.whole_content())

def extract_table(path):
    document = Document(path)
    for table in document.tables:
        data = [[cell.text for cell in row.cells] for row in table.rows]
        df = pd.DataFrame(data)
        df = df.rename(columns=df.iloc[0]).drop(df.index[0]).reset_index(drop=True)
        #print(df)

def extract_elements(path):
    document = Document(file_path)
    print(dir(document))
    content = ""
    for element in document.iter_inner_content():
        #print(element)
        if type(element) == Paragraph:
            content += element.text + '\n\n'
        elif type(element) == Table:
            data = [[cell.text for cell in row.cells] for row in element.rows]
            df = pd.DataFrame(data)
            df = df.rename(columns=df.iloc[0]).drop(df.index[0]).reset_index(drop=True)
            content += df.to_string(index=False)
    print(content)

def extract_whole_doc(path):
    loader = UnstructuredWordDocumentLoader(path)
    data = loader.load()
    print(data[0].page_content)


if __name__ == "__main__":
    #create_test_input(file_path)
    #extract_whole_doc(file_path)
    extract_elements(file_path)